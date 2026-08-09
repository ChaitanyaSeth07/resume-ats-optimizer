"""
Report Builder
Creates a presentable analysis report (DOCX) from evaluation metrics.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import pandas as pd
from typing import List, Dict


def _set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def build_analysis_report(logs: List[Dict]) -> BytesIO:
    """
    Build a clean, presentable DOCX analysis report from metric logs.
    """
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    df = pd.DataFrame(logs)

    # Convert numeric columns safely
    for col in ["overall_score", "keyword_score", "structural_score", "attempts_used", "target_score"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    # ---------- Title ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Resume ATS Optimizer")
    _set_run(run, size=20, bold=True, color=(15, 23, 42))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Creator Analysis Report")
    _set_run(run, size=14, bold=True, color=(37, 99, 235))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _set_run(run, size=10, color=(100, 116, 139))

    doc.add_paragraph("")

    # ---------- Summary ----------
    h = doc.add_paragraph()
    run = h.add_run("1. Summary Metrics")
    _set_run(run, size=13, bold=True, color=(15, 23, 42))

    total = len(df)
    avg_overall = df["overall_score"].mean() if total else 0
    avg_keyword = df["keyword_score"].mean() if total else 0
    avg_structure = df["structural_score"].mean() if total else 0
    avg_attempts = df["attempts_used"].mean() if total else 0

    summary_lines = [
        f"Total Evaluations: {total}",
        f"Average Overall Score: {avg_overall:.1f} / 100",
        f"Average Keyword Match: {avg_keyword:.1f}%",
        f"Average Structure Score: {avg_structure:.1f}%",
        f"Average Attempts Used: {avg_attempts:.1f}",
    ]

    for line in summary_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            _set_run(run, size=11, color=(51, 65, 85))

    doc.add_paragraph("")

    # ---------- Rating Distribution ----------
    h = doc.add_paragraph()
    run = h.add_run("2. Rating Distribution")
    _set_run(run, size=13, bold=True, color=(15, 23, 42))

    if "rating" in df.columns and total > 0:
        rating_counts = df["rating"].value_counts()
        for rating, count in rating_counts.items():
            p = doc.add_paragraph(f"{rating}: {count}")
            for run in p.runs:
                _set_run(run, size=11, color=(51, 65, 85))
    else:
        p = doc.add_paragraph("No rating data available.")
        for run in p.runs:
            _set_run(run, size=11, color=(100, 116, 139))

    doc.add_paragraph("")

    # ---------- Attempts Distribution ----------
    h = doc.add_paragraph()
    run = h.add_run("3. Feedback Attempts Distribution")
    _set_run(run, size=13, bold=True, color=(15, 23, 42))

    if "attempts_used" in df.columns and total > 0:
        attempt_counts = df["attempts_used"].value_counts().sort_index()
        for attempt, count in attempt_counts.items():
            p = doc.add_paragraph(f"{int(attempt)} attempt(s): {count} run(s)")
            for run in p.runs:
                _set_run(run, size=11, color=(51, 65, 85))
    else:
        p = doc.add_paragraph("No attempt data available.")
        for run in p.runs:
            _set_run(run, size=11, color=(100, 116, 139))

    doc.add_paragraph("")

    # ---------- Detailed Table ----------
    h = doc.add_paragraph()
    run = h.add_run("4. Detailed Evaluation Log")
    _set_run(run, size=13, bold=True, color=(15, 23, 42))

    if total == 0:
        p = doc.add_paragraph("No evaluations logged yet.")
        for run in p.runs:
            _set_run(run, size=11, color=(100, 116, 139))
    else:
        # Select and order useful columns
        columns = [
            c for c in [
                "timestamp",
                "overall_score",
                "keyword_score",
                "structural_score",
                "rating",
                "attempts_used",
                "target_score",
                "target_reached"
            ] if c in df.columns
        ]

        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"

        # Header
        header_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            header_cells[i].text = col.replace("_", " ").title()
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    _set_run(run, size=9, bold=True, color=(15, 23, 42))

        # Rows
        for _, row in df[columns].iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(columns):
                value = row[col]
                if pd.isna(value):
                    cells[i].text = ""
                else:
                    cells[i].text = str(value)
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        _set_run(run, size=9, color=(51, 65, 85))

    doc.add_paragraph("")

    # ---------- Notes ----------
    h = doc.add_paragraph()
    run = h.add_run("5. Notes")
    _set_run(run, size=13, bold=True, color=(15, 23, 42))

    notes = [
        "This report contains metrics only. Resume content and job descriptions are not stored.",
        "Overall Score combines keyword match and structural quality.",
        "Attempts represent how many feedback-loop cycles were used.",
        "Use this report for research evaluation, portfolio documentation, and iteration analysis."
    ]

    for note in notes:
        p = doc.add_paragraph(f"• {note}")
        for run in p.runs:
            _set_run(run, size=10, color=(51, 65, 85))

    # Save to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer