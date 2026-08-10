"""
DOCX Builder with user-selectable design options (ATS-safe)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re
from typing import Dict


THEMES = {
    "Blue": (37, 99, 235),
    "Charcoal": (31, 41, 55),
    "Teal": (13, 148, 136),
    "Green": (22, 163, 74),
    "Burgundy": (153, 27, 27),
}

FONTS = {
    "Calibri": "Calibri",
    "Arial": "Arial",
    "Georgia": "Georgia",
    "Garamond": "Garamond",
}

SPACING = {
    "Compact": {"before": 6, "after": 2, "body_after": 2, "line": 1.08},
    "Normal": {"before": 10, "after": 4, "body_after": 3, "line": 1.15},
    "Comfortable": {"before": 14, "after": 6, "body_after": 4, "line": 1.2},
}


def _set_run(run, font="Calibri", size=11, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_paragraph(paragraph, before=0, after=6, line=1.15, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def _add_bottom_border(paragraph, color_hex="2563EB", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _rgb_to_hex(rgb):
    return "{:02X}{:02X}{:02X}".format(*rgb)


def parse_optimized_resume(text: str) -> dict:
    sections = {
        "contact": "",
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "certifications": "",
        "other": "",
    }
    current = "other"
    for line in text.split("\n"):
        upper = line.strip().upper()
        if upper.startswith("CONTACT"):
            current = "contact"; continue
        if upper.startswith("SUMMARY"):
            current = "summary"; continue
        if upper.startswith("EXPERIENCE"):
            current = "experience"; continue
        if upper.startswith("EDUCATION"):
            current = "education"; continue
        if upper.startswith("SKILLS"):
            current = "skills"; continue
        if upper.startswith("PROJECTS"):
            current = "projects"; continue
        if upper.startswith("CERTIFICATIONS"):
            current = "certifications"; continue
        if upper.startswith("OTHER"):
            current = "other"; continue
        if line.strip():
            sections[current] += line + "\n"
    return sections


def create_ats_docx(optimized_text: str, design: Dict = None) -> BytesIO:
    design = design or {}
    theme_name = design.get("theme", "Blue")
    font_name = FONTS.get(design.get("font", "Calibri"), "Calibri")
    header_style = design.get("header_style", "Centered")
    section_style = design.get("section_style", "Underline")
    spacing_name = design.get("spacing", "Normal")
    accent_strength = design.get("accent_strength", "Medium")

    accent = THEMES.get(theme_name, THEMES["Blue"])
    accent_hex = _rgb_to_hex(accent)
    space = SPACING.get(spacing_name, SPACING["Normal"])
    border_size = "12" if accent_strength == "Medium" else "6"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)

    parsed = parse_optimized_resume(optimized_text)

    # ---------- Header / Contact ----------
    contact_lines = [l.strip() for l in parsed["contact"].strip().split("\n") if l.strip()]
    name = ""
    contact_parts = []

    for i, line in enumerate(contact_lines):
        cleaned = re.sub(r"^(Name|Email|Phone|LinkedIn|Location):\s*", "", line, flags=re.IGNORECASE).strip()
        if i == 0 or line.lower().startswith("name"):
            name = cleaned
        elif cleaned:
            contact_parts.append(cleaned)

    align = WD_ALIGN_PARAGRAPH.CENTER if header_style == "Centered" else WD_ALIGN_PARAGRAPH.LEFT

    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        _set_run(run, font=font_name, size=18 if header_style != "Minimal" else 16, bold=True, color=accent if header_style != "Minimal" else (15, 23, 42))
        _set_paragraph(p, before=0, after=2, align=align)

    if contact_parts:
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(contact_parts))
        _set_run(run, font=font_name, size=9.5, color=(51, 65, 85))
        _set_paragraph(p, before=0, after=6, align=align)

    # Header line
    line_p = doc.add_paragraph()
    _set_paragraph(line_p, before=0, after=8)
    _add_bottom_border(line_p, color_hex=accent_hex, size=border_size)

    def add_section(title: str, content: str):
        if not content.strip():
            return

        h = doc.add_paragraph()
        run = h.add_run(title.upper())
        _set_run(run, font=font_name, size=11.5, bold=True, color=accent if section_style != "Simple bold" else (15, 23, 42))
        _set_paragraph(h, before=space["before"], after=space["after"])

        if section_style == "Underline":
            _add_bottom_border(h, color_hex=accent_hex, size="6")
        elif section_style == "Caps + line":
            _add_bottom_border(h, color_hex="94A3B8", size="4")

        for raw in content.strip().split("\n"):
            line = raw.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if line.startswith(("-", "•", "*", "–", "—")):
                clean = re.sub(r"^[-•*–—]\s*", "", line)
                run = p.add_run("• " + clean)
            else:
                run = p.add_run(line)
            _set_run(run, font=font_name, size=10.5, color=(30, 41, 59))
            _set_paragraph(p, before=0, after=space["body_after"], line=space["line"])

    add_section("Professional Summary", parsed["summary"])
    add_section("Experience", parsed["experience"])
    add_section("Education", parsed["education"])
    add_section("Skills", parsed["skills"])
    add_section("Projects", parsed["projects"])
    add_section("Certifications", parsed["certifications"])
    if parsed["other"].strip():
        add_section("Additional Information", parsed["other"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer